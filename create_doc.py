from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── Page margins ───
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── Style helpers ───
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

for level, size, color in [(1, 26, '1F4E79'), (2, 18, '2E75B6'), (3, 14, '2E75B6')]:
    s = doc.styles[f'Heading {level}']
    s.font.size = Pt(size)
    s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = True
    s.font.name = 'Calibri'
    s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    s.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    s.paragraph_format.space_after = Pt(8)

def add_shaded_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F4E79')
        cell._tc.get_or_add_tcPr().append(shading)
    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r_idx % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'E8F0FE')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    return p

# ═══════════════════════════════════════════════════════════════
# COVER / TITLE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('KiidScreen')
run.font.size = Pt(44)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('نظام حماية الأطفال من المحتوى الهابط\nتحت إشراف الوالدين')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('مسودة خطة المشروع الأولية\n\n').font.size = Pt(13)
meta.add_run('تاريخ المسودة: ').font.size = Pt(11)
meta.add_run('11 مايو 2026').font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════
doc.add_heading('فهرس المحتويات', level=1)
toc_items = [
    '1. نظرة عامة عن المشروع',
    '2. أهداف النظام',
    '3. الجمهور المستهدف',
    '4. هيكل المشروع (الملفات والمجلدات)',
    '5. تدفق العمل (User Flow)',
    '  5.1 وضع الطفل (الشاشة الرئيسية)',
    '  5.2 الدخول للوحة تحكم الوالدين',
    '  5.3 لوحة تحكم الوالدين',
    '6. تفاصيل تقنية',
    '  6.1 تشغيل الفيديو بدون اقتراحات',
    '  6.2 كشف نوع رابط الفيديو',
    '  6.3 إدارة الحالة والتخزين (Firebase + محلي)',
    '  6.4 حماية الخروج',
    '  6.5 عداد الوقت',
    '  6.6 التصميم المتجاوب والتطبيق الجوال',
    '7. هيكل Firebase (Firestore)',
    '8. خطة التنفيذ (المراحل والجداول الزمنية)',
    '9. الملحقات والمراجع',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)
        if not item.startswith('  '):
            run.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. نظرة عامة
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. نظرة عامة عن المشروع', level=1)
add_para(doc, 'KiidScreen هو تطبيق ويب متكامل يهدف إلى حماية الأطفال من المحتوى الهابط والغير مناسب على الإنترنت، مع إعطاء أولياء الأمور تحكماً كاملاً بما يشاهده أطفالهم. يعمل التطبيق حالياً كصفحة ويب متصفح ، وهو مُهيأ ليتحول لاحقاً إلى تطبيق جوال (Android / iOS).')
add_para(doc, 'الفكرة الأساسية: صفحة مقفلة (Kiosk Mode) تعرض فيديوهات يختارها الوالدان، لا يمكن للطفل الخروج منها إلا بطريقة يحددها الوالدان (رمز PIN سري + عدد معين من الضغطات على زر الرجوع).')

# ═══════════════════════════════════════════════════════════════
# 2. أهداف النظام
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. أهداف النظام', level=1)
goals = [
    'توفير بيئة مشاهدة آمنة للأطفال دون التعرض لمحتوى عشوائي أو غير مناسب.',
    'تمكين الوالدين من إدارة كامل المحتوى المعروض (إضافة، حذف، ترتيب فيديوهات).',
    'دعم مصادر متعددة للفيديو: YouTube (بدون اقتراحات)، Vimeo، وروابط مباشرة.',
    'التحكم بوقت المشاهدة عبر عداد زمني، مع إطفاء الشاشة تلقائياً عند انتهاء الوقت.',
    'حماية قوية ضد محاولات الخروج من التطبيق من قبل الطفل.',
    'تخزين سحابي عبر Firebase يسمح بمزامنة الإعدادات عبر الأجهزة.',
    'جاهزية للتحويل إلى تطبيق جوال (معمارية متجاوبة + PWA).',
]
for g in goals:
    add_bullet(doc, g)

# ═══════════════════════════════════════════════════════════════
# 3. الجمهور المستهدف
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. الجمهور المستهدف', level=1)
add_para(doc, 'أولياء أمور يرغبون في توفير بيئة مشاهدة آمنة لأطفالهم (من سن 3 إلى 12 سنة)، مع سهولة الاستخدام وعدم الحاجة لخبرة تقنية متقدمة.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. هيكل المشروع
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. هيكل المشروع (الملفات والمجلدات)', level=1)
add_para(doc, 'الملفات التالية تشكل المشروع الكامل. جميع الملفات تعمل معاً لتوفير تجربة متكاملة:')

tree_rows = [
    ['kiidscreen/', '', 'المجلد الرئيسي للمشروع'],
    ['', 'index.html', 'الصفحة الرئيسية - التوجيه بين وضع الطفل ولوحة التحكم'],
    ['', 'manifest.json', 'ملف PWA لتثبيت التطبيق على شاشة الجهاز'],
    ['', 'sw.js', 'Service Worker - دعم التشغيل بدون إنترنت'],
    ['css/', '', 'ملفات الأنماط'],
    ['', '  style.css', 'جميع الأنماط (Kid mode + Dashboard + Responsive)'],
    ['js/', '', 'ملفات الجافاسكريبت'],
    ['', '  app.js', 'التوجيه الرئيسي - التبديل بين وضع الطفل ولوحة التحكم'],
    ['', '  video-handler.js', 'كشف نوع رابط الفيديو وتشغيل المشغل المناسب'],
    ['', '  timer.js', 'عداد الوقت (ساعات ودقائق) مع تخزين الحالة ومقاومة إعادة التحميل'],
    ['', '  lock.js', 'حماية الخروج: PIN + عدّ ضغطات زر الرجوع'],
    ['', '  dashboard.js', 'لوحة تحكم الوالدين بالكامل (CRUD فيديوهات، إعدادات)'],
    ['', '  firebase-init.js', 'تهيئة Firebase: Firebase Auth + Firestore'],
    ['firebase.json', '', 'إعدادات مشروع Firebase'],
    ['create_doc.py', '', 'سكريبت إنشاء ملف التوثيق (هذا الملف)'],
]
add_shaded_table(doc, ['المجلد', 'الملف', 'الوصف'], tree_rows, col_widths=[4, 5, 8])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. تدفق العمل
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. تدفق العمل (User Flow)', level=1)

doc.add_heading('5.1 وضع الطفل (الشاشة الرئيسية)', level=2)
add_para(doc, 'عند فتح التطبيق، يدخل الطفل مباشرة إلى وضع المشاهدة الآمن. التدفق كالتالي:')

steps = [
    ('فتح التطبيق / تحميل الصفحة ← ', 'الدخول المباشر لوضع الطفل دون أي خطوات إضافية.'),
    ('شاشة كاملة إجبارية ← ', 'تفعيل Fullscreen API تلقائياً. يمنع تصغير المتصفح أو رؤية شريط العنوان.'),
    ('تشغيل قائمة الفيديو ← ', 'تبدأ playlist تلقائياً بأول فيديو في القائمة. تنتقل للفيديو التالي عند انتهاء الفيديو الحالي.'),
    ('YouTube بدون اقتراحات ← ', 'يتم تضمين الفيديو باستخدام IFrame API مع المعاملات: rel=0 (لا اقتراحات)، modestbranding=1 (علامة مائية مخففة)، controls=1 مع إخفاء recommended videos.'),
    ('عداد الوقت ← ', 'يظهر بشكل غير مزعج (دائرة تقدم صغيرة في الزاوية). يحسب الوقت المتبقي.'),
    ('انتهاء الوقت ← ', 'تتحول الشاشة إلى اللون الأسود بالكامل + إيقاف الفيديو + ظهور رسالة "انتهى وقت المشاهدة" + طلب PIN من الوالدين لإعادة التشغيل.'),
    ('محاولة الخروج ← ', 'الضغط على زر الرجوع في المتصفح: يزيد العداد الداخلي. عند الوصول للعدد المطلوب (configurable) → ظهور overlay إدخال PIN.'),
]
for bold, normal in steps:
    add_bullet(doc, normal, bold_prefix=bold)

doc.add_heading('5.2 الدخول للوحة تحكم الوالدين', level=2)
add_para(doc, 'هناك طريقتان للوصول للوحة التحكم:')
methods = [
    ('الطريقة الأولى: أيقونة مخفية ← ', 'النقر على أيقونة صغيرة أو النقر 5 مرات متتالية على الزاوية اليمنى السفلى من الشاشة. تظهر شاشة إدخال PIN.'),
    ('الطريقة الثانية: من شاشة PIN ← ', 'عند ظهور شاشة PIN (بعد انتهاء الوقت أو بعد محاولة الخروج)، يوجد زر إضافي "دخول للوحة التحكم".'),
]
for bold, normal in methods:
    add_bullet(doc, normal, bold_prefix=bold)

doc.add_heading('5.3 لوحة تحكم الوالدين', level=2)
add_para(doc, 'بعد إدخال PIN الصحيح، تظهر لوحة التحكم مقسمة إلى الأقسام التالية:')

sections = [
    ('إدارة الفيديو: ', 'إضافة فيديو جديد (YouTube URL / Vimeo URL / Direct URL)، ترتيب الفيديو (سحب وإفلات)، حذف فيديو، معاينة الفيديو قبل الإضافة.'),
    ('إعدادات الوقت: ', 'تحديد عدد الساعات (0-12)، تحديد عدد الدقائق (0-59).'),
    ('إعدادات الحماية: ', 'عدد ضغطات الرجوع المطلوبة للخروج (1-10)، تغيير PIN (4-6 أرقام مع تأكيد)، تفعيل/تعطيل PIN.'),
    ('إعدادات Firebase: ', 'رابط مشروع Firebase، مفتاح API (Web API Key).'),
    ('معاينة وضع الطفل: ', 'زر لرؤية كيف سيبدو التطبيق للطفل قبل الخروج من لوحة التحكم.'),
]
for bold, normal in sections:
    add_bullet(doc, normal, bold_prefix=bold)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. تفاصيل تقنية
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. تفاصيل تقنية', level=1)

doc.add_heading('6.1 تشغيل الفيديو بدون اقتراحات', level=2)
add_para(doc, 'لضمان عدم ظهور أي محتوى غير مرغوب للطفل، يتم استخدام الطرق التالية لكل منصة فيديو:')

video_rows = [
    ['YouTube', '<iframe> مع YouTube IFrame API', 'rel=0 (إخفاء الاقتراحات)\nmodestbranding=1 (علامة مائية مخففة)\nenablejsapi=1 (تحكم برمجي)\nplaysinline=1 (تشغيل داخلي)\ncc_load_policy=0 (إخفاء الترجمة)', 'نعم'],
    ['Vimeo', '<iframe> مع Vimeo Player API', 'byline=0 (إخفاء اسم الكاتب)\ntitle=0 (إخفاء العنوان)\nportrait=0 (إخفاء الصورة)\ndnt=1 (Do Not Track)', 'نعم'],
    ['رابط مباشر', '<video> عنصر HTML5', 'controls (عناصر التحكم)\nيتحقق من الامتداد: mp4, webm, ogg', 'يدوياً'],
    ['رابط آخر', 'حاول embed + fallback', 'إن كان الرابط منصة غير معروفة، يحاول تضمينه بصيغة embed عامة، وإلا يظهر رسالة خطأ', '-'],
]
add_shaded_table(doc, ['المصدر', 'الطريقة', 'المعاملات', 'قابلية التحكم البرمجي'], video_rows, col_widths=[2.5, 4, 7, 3])

doc.add_heading('6.2 كشف نوع رابط الفيديو', level=2)
add_para(doc, 'دالة الكشف (auto-detect) تفحص الرابط المُدخل وتقرر أي مشغل تستخدم:')

detect_rows = [
    ['youtube.com/watch?v=', 'YouTube', 'استخراج الـ Video ID من معامل v='],
    ['youtu.be/', 'YouTube', 'استخراج الـ Video ID من المسار'],
    ['youtube.com/embed/', 'YouTube', 'استخراج الـ Video ID مباشر'],
    ['vimeo.com/', 'Vimeo', 'استخراج الـ Video ID من المسار'],
    ['ينتهي بـ .mp4', 'مباشر (HTML5)', 'تضمين <video> مع المسار'],
    ['ينتهي بـ .webm', 'مباشر (HTML5)', 'تضمين <video> مع المسار'],
    ['ينتهي بـ .ogg', 'مباشر (HTML5)', 'تضمين <video> مع المسار'],
    ['أي رابط آخر', 'محاولة تضمين', 'حاول <iframe> مباشر مع fallback لرسالة خطأ'],
]
add_shaded_table(doc, ['النمط في الرابط', 'النوع المُكتشف', 'آلية المعالجة'], detect_rows, col_widths=[5, 3, 8])

doc.add_heading('6.3 إدارة الحالة والتخزين', level=2)
add_para(doc, 'لضمان استمرارية البيانات عبر الجلسات، يتم استخدام ثلاث طبقات للتخزين:')

storage_rows = [
    ['Firebase Firestore', 'سحابي', 'قائمة الفيديو، إعدادات PIN، إعدادات الوقت، إعدادات الحماية', 'تخزين دائم، مزامنة عبر الأجهزة، استرجاع عند فتح التطبيق من أي متصفح/جهاز'],
    ['localStorage', 'محلي', 'نسخة مؤقتة من PIN (لصرف أقل عدد من الطلبات لل Firebase)', 'تسريع الدخول للوحة التحكم، يعمل بدون إنترنت'],
    ['sessionStorage', 'محلي (جلسة)', 'حالة عداد الوقت (وقت البدء، الوقت المتبقي)', 'مقاومة إعادة تحميل الصفحة (F5) - لا يفقد العداد حالته'],
]
add_shaded_table(doc, ['الطبقة', 'النوع', 'البيانات المخزنة', 'السبب'], storage_rows, col_widths=[3, 1.5, 6, 6])

doc.add_heading('6.4 حماية الخروج', level=2)
add_para(doc, 'تمنع حماية الخروج الطفل من مغادرة التطبيق باتباع آلية متعددة الطبقات:')
exit_rows = [
    ['1. مستمع popstate', 'يكتشف الضغط على زر الرجوع في المتصفح',
     'يتم دفع حالة جديدة إلى history في كل مرة لمنع الخروج الفعلي، ويزداد العداد الداخلي.'],
    ['2. عداد ضغطات الرجوع', 'بعد N ضغطة → PIN overlay',
     'العدد N قابل للتكوين من لوحة التحكم (1-10). يتم إعادة تعيين العداد (Reset) بعد 3 ثوانٍ من عدم الضغط على رجوع.'],
    ['3. PIN overlay', 'نموذج إدخال PIN (4-6 أرقام)',
     'عند الوصول للحد المطلوب من الضغطات، تظهر شاشة كاملة تطلب PIN. يتم التحقق عبر Firebase.'],
    ['4. OnBeforeUnload', 'يمنع إغلاق المتصفح',
     'يعرض تأكيد الخروج. في حال تم الإغلاق، عند العودة يعود لوضع الطفل مباشرة إذا كانت الجلسة نشطة.'],
    ['5. Fullscreen lock', 'يمنع تصغير الشاشة',
     'يستمع لحدث FullscreenChange ويعيد تفعيل ملء الشاشة تلقائياً عند محاولة الخروج منها.'],
]
add_shaded_table(doc, ['الطبقة', 'الوصف', 'التفاصيل'], exit_rows, col_widths=[3.5, 4, 9])

doc.add_heading('6.5 عداد الوقت', level=2)
add_para(doc, 'نظام إدارة وقت المشاهدة:')
timer_rows = [
    ['تحديد الوقت', 'من لوحة التحكم', 'ساعات (0-12) + دقائق (0-59)'],
    ['الحساب', 'برمجياً', 'تحويل الوقت المحدد إلى ثواني وحساب الفرق مع وقت البدء'],
    ['التخزين', 'sessionStorage', 'تخزين وقت البدء (startTime) لحساب الوقت المتبقي بدقة حتى مع إعادة التحميل'],
    ['عرض التقدم', 'دائرة تقدم (CSS Progress Ring)', 'تظهر بشكل صغير غير مزعج في الزاوية السفلى اليمنى ± نص "X:XX متبقي"'],
    ['نهاية الوقت', 'حدث (Event)', 'إيقاف الفيديو + شاشة سوداء كاملة (CSS overlay) + PIN'],
    ['مقاومة التلاعب', 'تحقق دوري', 'فحص وقت البدء في كل ثانية ومقارنته بالوقت الحالي'],
]
add_shaded_table(doc, ['الميزة', 'الآلية', 'التفاصيل'], timer_rows, col_widths=[3, 3.5, 10])

doc.add_heading('6.6 التصميم المتجاوب والتطبيق الجوال', level=2)
add_para(doc, 'لضمان مستقبل التطبيق كتطبيق جوال، تم تضمين الممارسات التالية:')
pwa_rows = [
    ['Responsive CSS', '@media queries', 'أحجام خطوط نسبية (rem, em)، شبكات CSS (Flexbox/Grid)، تحجيم الفيديو تلقائياً'],
    ['PWA Manifest', 'manifest.json', 'اسم التطبيق، الأيقونات (48px-512px)، لون السمة، وضع العرض (standalone لإخفاء شريط المتصفح)'],
    ['Service Worker', 'sw.js', 'تخزين مؤقت للملفات الأساسية (HTML, CSS, JS) لتشغيل التطبيق بدون إنترنت أو مع ضعف الشبكة'],
    ['Touch Events', 'معالجة اللمس', 'الأزرار كبيرة الحجم (≥ 48px)، دعم اللمس في السحب والإفلات، دوار بأحجام مناسبة'],
    ['جاهزية للتحويل', 'Native Wrapper', 'المعمارية تسمح بالتغليف باستخدام Capacitor أو Cordova أو React Native WebView دون تغيير الكود الأساسي'],
]
add_shaded_table(doc, ['الميزة', 'التقنية', 'التفاصيل'], pwa_rows, col_widths=[3.5, 4, 9])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. هيكل Firebase
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. هيكل Firebase (Firestore)', level=1)
add_para(doc, 'بيانات التطبيق مخزنة في Firestore وفق الهيكل التالي:')

firebase_rows = [
    ['settings', '{parentId}', 'pin', 'string', 'رمز PIN (SHA-256 hashed)'],
    ['', '', 'backPressCount', 'number', 'عدد ضغطات الرجوع المطلوبة (default: 3)'],
    ['', '', 'timeLimitHours', 'number', 'الحد الأقصى للساعات (default: 1)'],
    ['', '', 'timeLimitMinutes', 'number', 'الحد الأقصى للدقائق (default: 0)'],
    ['', '', 'fullscreen', 'boolean', 'تفعيل ملء الشاشة (default: true)'],
    ['', '', 'createdAt', 'timestamp', 'تاريخ الإنشاء'],
    ['', '', 'updatedAt', 'timestamp', 'آخر تحديث'],
    ['videos', '{parentId}', '-', 'collection', 'مجموعة الفيديو'],
    ['', 'list/{videoId}', 'url', 'string', 'رابط الفيديو'],
    ['', '', 'type', 'string', 'youtube / vimeo / direct'],
    ['', '', 'title', 'string', 'عنوان الفيديو (اختياري)'],
    ['', '', 'order', 'number', 'ترتيب العرض (للـ playlist)'],
    ['', '', 'addedAt', 'timestamp', 'تاريخ الإضافة'],
]
add_shaded_table(doc, ['Collection', 'المستند', 'الحقل', 'النوع', 'الوصف'], firebase_rows, col_widths=[2.5, 3, 3.5, 2, 5.5])

doc.add_paragraph()
add_para(doc, 'ملاحظة أمنية: الـ PIN يُخزّن بشكل مشفر (SHA-256) وليس كنص صريح. المقارنة تتم على التشفير وليس على النص الأصلي.', bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. خطة التنفيذ
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. خطة التنفيذ (المراحل والجداول الزمنية)', level=1)
add_para(doc, 'المشروع مقسم إلى 7 مراحل، كل مرحلة تبني على سابقتها:')

phase_rows = [
    ['1', 'إعداد الهيكل', 'يوم',
     'إنشاء المجلدات والملفات الأساسية\nتهيئة Firebase (المشروع + Auth + Firestore)\nالصفحة الرئيسية مع التوجيه (hash-based routing)\nملف الأنماط الأساسي'],
    ['2', 'مشغل الفيديو', 'يوم',
     'دالة كشف نوع الرابط (YouTube / Vimeo / Direct)\nYouTube IFrame API مع rel=0\nVimeo IFrame API\nHTML5 video للمباشر\nتشغيل تلقائي لل playlist'],
    ['3', 'عداد الوقت', 'يوم',
     'واجهة إعداد الوقت\nحساب الوقت المتبقي\nدائرة التقدم\nتخزين الحالة في sessionStorage\nحدث نهاية الوقت (شاشة سوداء + إيقاف)'],
    ['4', 'حماية الخروج', 'يوم',
     'مستمع popstate + عدّ الضغطات\nإعادة تعيين العداد\nPIN overlay\nFullscreen API\nمنع الإغلاق'],
    ['5', 'لوحة التحكم', 'يوم',
     'واجهة إدارة الفيديو (إضافة/حذف/ترتيب)\nإعدادات الوقت\nإعدادات PIN وضغطات الرجوع\nالاتصال بـ Firebase (قراءة/كتابة)\nأمان المسار (تحقق PIN)'],
    ['6', 'PWA + تصميم متجاوب', 'يوم',
     'manifest.json مع الأيقونات\nService Worker للتخزين المؤقت\nCSS Responsive (جوال، لوحي، مكتبي)\nتحسينات اللمس\nاختبار التوافق'],
    ['7', 'اختبارات وتحسينات', 'يوم',
     'اختبار شامل لجميع المسارات\nتحسين الأداء\nتجربة المستخدم (UX)\nإصلاح الأخطاء\nتوثيق الاستخدام'],
]
add_shaded_table(doc, ['المرحلة', 'الاسم', 'المدة', 'المهام'], phase_rows, col_widths=[1.5, 3, 1.5, 10.5])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. الملحقات
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. الملحقات والمراجع', level=1)

doc.add_heading('9.1 التقنيات المستخدمة', level=2)
tech_rows = [
    ['الواجهة الأمامية', 'HTML5, CSS3, JavaScript (ES6+)'],
    ['التخزين السحابي', 'Firebase Firestore + Firebase Auth'],
    ['YouTube API', 'YouTube IFrame Player API (developers.google.com/youtube/iframe_api_reference)'],
    ['Vimeo API', 'Vimeo Player API (developer.vimeo.com/player/sdk)'],
    ['PWA', 'Web App Manifest + Service Worker API'],
    ['التغليف للموبايل (مستقبلاً)', 'Capacitor (ionicframework.com/capacitor) أو React Native WebView'],
]
add_shaded_table(doc, ['المكون', 'التقنية'], tech_rows, col_widths=[5, 11.5])

doc.add_heading('9.2 ملاحظات إضافية', level=2)
notes = [
    'النسخة الحالية تعمل على المتصفحات الحديثة (Chrome, Firefox, Edge, Safari).',
    'دعم Internet Explorer غير مطلوب.',
    'خدمة Firebase لها طبقة مجانية (Free Tier) تكفي للمشاريع الصغيرة والمتوسطة: 1GB تخزين، 10GB تحميل/شهرياً، 50K قراءة/يوم.',
    'لتحويل التطبيق لجوال مستقبلاً، يمكن استخدام Capacitor دون تغيير كود الواجهة.',
    'الـ PIN يُخزّن مشفراً باستخدام SHA-256 للمقارنة الآمنة.',
]
for n in notes:
    add_bullet(doc, n)

# ─── Footer ───
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— نهاية المسودة —')
run.font.size = Pt(14)
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

# ─── Save ───
output_path = os.path.join(os.path.dirname(__file__), 'KiidScreen_Project_Plan.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
